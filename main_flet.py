import flet as ft
from flet import Page, TextField, ElevatedButton, Column, Row, Container, Text, Icon, Icons, AlertDialog, TextButton, AppBar, IconButton, ListView, Divider
import sqlite3
import json
import os
import tempfile
import webbrowser
import subprocess
import platform
import asyncio
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- 路径处理（移动端使用应用文档目录）----------
def get_base_path():
    if platform.system() == "Android" or platform.system() == "iOS":
        return os.path.join(os.environ.get('HOME', ''), 'documents')
    else:
        return os.path.dirname(__file__)

BASE_PATH = get_base_path()
DB_PATH = os.path.join(BASE_PATH, "legal_app.db")
CONFIG_FILE = os.path.join(BASE_PATH, "config.ini")

# ---------- 数据库初始化 ----------
def init_db():
    try:
        os.makedirs(BASE_PATH, exist_ok=True)
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            phone TEXT UNIQUE,
            username TEXT UNIQUE,
            password TEXT,
            is_purchased INTEGER DEFAULT 0
        )''')
        c.execute('''CREATE TABLE IF NOT EXISTS templates (
            id INTEGER PRIMARY KEY,
            name TEXT,
            category TEXT,
            fields TEXT,
            content_template TEXT
        )''')
        c.execute('''CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            template_id INTEGER,
            filled_data TEXT,
            generated_content TEXT,
            create_time TEXT
        )''')
        c.execute('''CREATE TABLE IF NOT EXISTS trial (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            install_date TEXT,
            last_run_date TEXT
        )''')
        conn.commit()

        c.execute("SELECT COUNT(*) FROM templates")
        if c.fetchone()[0] == 0:
            # ========== 完整模板列表（23个）==========
            templates = [
                (1, "用工合同", "合同类",
                 '[{"name":"用人单位","hint":"例如：某某科技有限公司","required":true},{"name":"劳动者","hint":"例如：张三","required":true},{"name":"工作岗位","hint":"例如：软件工程师","required":true},{"name":"合同期限","hint":"例如：2025年4月1日至2028年3月31日","required":true},{"name":"月工资","hint":"例如：8000元","required":true}]',
                 "劳动合同\n\n甲方（用人单位）：{{用人单位}}\n乙方（劳动者）：{{劳动者}}\n\n根据《中华人民共和国劳动法》，甲乙双方经平等协商，自愿签订本合同。\n\n第一条 合同期限\n本合同期限为{{合同期限}}。\n\n第二条 工作岗位\n乙方同意在{{工作岗位}}岗位工作。\n\n第三条 劳动报酬\n甲方每月以货币形式支付乙方工资，月工资为{{月工资}}。\n\n第四条 社会保险\n甲方按国家规定为乙方缴纳社会保险。\n\n第五条 违约责任\n任何一方违反本合同，应承担相应法律责任。\n\n第六条 争议解决\n双方发生争议，可向劳动争议仲裁委员会申请仲裁。\n\n本合同一式两份，双方各执一份。\n\n甲方（盖章）：________    乙方（签字）：________\n日期：____年__月__日"),
                (2, "买卖合同", "合同类",
                 '[{"name":"卖方","hint":"例如：某某贸易公司","required":true},{"name":"买方","hint":"例如：某某工厂","required":true},{"name":"标的物","hint":"例如：钢材100吨","required":true},{"name":"数量","hint":"例如：100吨","required":true},{"name":"单价","hint":"例如：4000元/吨","required":true}]',
                 "买卖合同\n\n合同编号：________\n\n卖方：{{卖方}}\n买方：{{买方}}\n\n第一条 标的物、数量、价款\n标的物名称：{{标的物}}\n数量：{{数量}}\n单价：{{单价}}\n总价：{{数量}} * {{单价}} = ________元\n\n第二条 质量标准\n按国家标准执行。\n\n第三条 交货时间、地点\n卖方应于____年__月__日前将货物交付至________。\n\n第四条 付款方式\n买方应在收到货物后____日内付清全款。\n\n第五条 违约责任\n逾期交货或付款，每日按合同总价的0.1%支付违约金。\n\n第六条 争议解决\n提交买方所在地人民法院诉讼解决。\n\n本合同一式两份，双方各执一份。\n\n卖方（盖章）：________    买方（盖章）：________\n日期：____年__月__日"),
                (3, "销售合同", "合同类",
                 '[{"name":"销售方","hint":"例如：某某电子厂","required":true},{"name":"购买方","hint":"例如：某某经销商","required":true},{"name":"产品名称","hint":"例如：智能手机","required":true},{"name":"数量","hint":"例如：500台","required":true},{"name":"总金额","hint":"例如：500000元","required":true}]',
                 "销售合同\n\n销售方：{{销售方}}\n购买方：{{购买方}}\n\n第一条 产品规格\n产品名称：{{产品名称}}\n数量：{{数量}}\n总金额：{{总金额}}\n\n第二条 质量保证\n销售方保证产品符合国家质量标准。\n\n第三条 交货期\n销售方应于合同签订后____日内交货。\n\n第四条 付款方式\n购买方应于交货前支付50%货款，余款验收后7日内付清。\n\n第五条 售后服务\n保修期一年。\n\n第六条 争议解决\n提交销售方所在地法院管辖。\n\n销售方（盖章）：________    购买方（盖章）：________\n日期：____年__月__日"),
                (4, "采购合同", "合同类",
                 '[{"name":"采购方","hint":"例如：某某超市","required":true},{"name":"供应方","hint":"例如：某某食品厂","required":true},{"name":"产品名称","hint":"例如：速冻水饺","required":true},{"name":"总金额","hint":"例如：50000元","required":true}]',
                 "采购合同\n\n采购方：{{采购方}}\n供应方：{{供应方}}\n\n一、采购产品\n产品名称：{{产品名称}}\n总金额：{{总金额}}\n\n二、质量标准\n符合国家食品安全标准。\n\n三、交货时间\n供应方应于____年__月__日前交货。\n\n四、验收\n采购方在收货后3日内验收，不合格产品可退货。\n\n五、付款\n验收合格后15日内付清货款。\n\n六、违约责任\n逾期交货每日按合同总价的0.5%支付违约金。\n\n采购方（盖章）：________    供应方（盖章）：________\n日期：____年__月__日"),
                (5, "房屋租赁合同", "合同类",
                 '[{"name":"出租方","hint":"例如：李四","required":true},{"name":"承租方","hint":"例如：王五","required":true},{"name":"房屋地址","hint":"例如：北京市朝阳区某某小区1号楼101室","required":true},{"name":"租赁期限","hint":"例如：2025年5月1日至2026年4月30日","required":true},{"name":"月租金","hint":"例如：3000元","required":true}]',
                 "房屋租赁合同\n\n出租方（甲方）：{{出租方}}\n承租方（乙方）：{{承租方}}\n\n第一条 房屋基本情况\n房屋坐落：{{房屋地址}}\n建筑面积：________平方米\n\n第二条 租赁期限\n租赁期自{{租赁期限}}。\n\n第三条 租金及支付\n月租金为{{月租金}}。乙方应于每月5日前支付。\n\n第四条 押金\n乙方于签约时支付押金________元，租赁期满无违约退还。\n\n第五条 房屋使用\n乙方不得擅自改变房屋结构，不得转租。\n\n第六条 维修责任\n房屋自然损坏由甲方维修，人为损坏由乙方维修。\n\n第七条 违约责任\n任何一方提前解约，应支付一个月租金作为违约金。\n\n第八条 争议解决\n协商不成，可向房屋所在地法院起诉。\n\n本合同一式两份，双方各执一份。\n\n甲方（签字）：________    乙方（签字）：________\n日期：____年__月__日"),
                (6, "民事起诉状", "诉讼类",
                 '[{"name":"原告","hint":"例如：张三","required":true},{"name":"被告","hint":"例如：李四","required":true},{"name":"诉讼请求","hint":"例如：请求法院判令被告偿还借款10万元及利息","required":true},{"name":"事实与理由","hint":"例如：2024年1月，被告向原告借款10万元，约定2024年12月归还，但至今未还。原告多次催讨无果。","required":true}]',
                 "民事起诉状\n\n原告：{{原告}}，性别____，出生____年__月__日，住址________，身份证号________。\n被告：{{被告}}，性别____，出生____年__月__日，住址________，身份证号________。\n\n诉讼请求：\n{{诉讼请求}}\n\n事实与理由：\n{{事实与理由}}\n\n此致\n________人民法院\n\n具状人：________\n____年__月__日"),
                (7, "离婚起诉状", "诉讼类",
                 '[{"name":"原告","hint":"例如：王芳","required":true},{"name":"被告","hint":"例如：李强","required":true},{"name":"离婚请求","hint":"例如：请求判决离婚","required":true},{"name":"子女抚养安排","hint":"例如：儿子李小强由原告抚养，被告每月支付抚养费2000元","required":true},{"name":"财产分割请求","hint":"例如：婚后购买的房产归原告所有，原告补偿被告30万元","required":true}]',
                 "离婚起诉状\n\n原告：{{原告}}，女，____年__月__日生，住________。\n被告：{{被告}}，男，____年__月__日生，住________。\n\n诉讼请求：\n一、{{离婚请求}}；\n二、子女抚养：{{子女抚养安排}}；\n三、财产分割：{{财产分割请求}}。\n\n事实与理由：\n原被告于____年__月__日登记结婚，婚后生育一子。因性格不合，夫妻感情破裂，无和好可能。\n\n此致\n________人民法院\n\n具状人：________\n____年__月__日"),
                (8, "民事上诉状", "诉讼类",
                 '[{"name":"上诉人","hint":"例如：张三","required":true},{"name":"被上诉人","hint":"例如：李四","required":true},{"name":"上诉请求","hint":"例如：请求撤销一审判决，改判支持上诉人全部诉讼请求","required":true},{"name":"上诉理由","hint":"例如：一审认定事实错误，适用法律不当","required":true}]',
                 "民事上诉状\n\n上诉人：{{上诉人}}\n被上诉人：{{被上诉人}}\n\n上诉人因________纠纷一案，不服________人民法院____年__月__日（____）____号民事判决，现提起上诉。\n\n上诉请求：\n{{上诉请求}}\n\n上诉理由：\n{{上诉理由}}\n\n此致\n________中级人民法院\n\n上诉人：________\n____年__月__日"),
                (9, "刑事告诉状", "诉讼类",
                 '[{"name":"自诉人","hint":"例如：赵六","required":true},{"name":"被告人","hint":"例如：钱七","required":true},{"name":"罪名","hint":"例如：故意伤害罪","required":true},{"name":"事实与证据","hint":"例如：2025年1月，被告人在某某地点将自诉人打伤，医院诊断为轻伤一级。有监控录像和证人证言。","required":true}]',
                 "刑事自诉状\n\n自诉人：{{自诉人}}\n被告人：{{被告人}}\n\n案由：{{罪名}}\n\n诉讼请求：\n依法追究被告人刑事责任。\n\n事实与证据：\n{{事实与证据}}\n\n此致\n________人民法院\n\n自诉人：________\n____年__月__日"),
                (10, "刑事答辩状", "诉讼类",
                 '[{"name":"答辩人","hint":"例如：钱七","required":true},{"name":"答辩事由","hint":"例如：自诉人指控我故意伤害，但我当时是正当防卫","required":true}]',
                 "刑事答辩状\n\n答辩人：{{答辩人}}\n\n因自诉人________指控我________一案，现答辩如下：\n\n{{答辩事由}}\n\n此致\n________人民法院\n\n答辩人：________\n____年__月__日"),
                (11, "支付令申请书", "诉讼类",
                 '[{"name":"申请人","hint":"例如：某银行","required":true},{"name":"被申请人","hint":"例如：王某","required":true},{"name":"请求金额","hint":"例如：50000元","required":true},{"name":"事实理由","hint":"例如：被申请人于2024年3月向申请人借款5万元，约定2024年9月归还，至今未还。","required":true}]',
                 "支付令申请书\n\n申请人：{{申请人}}\n被申请人：{{被申请人}}\n\n请求事项：\n请求贵院发出支付令，督促被申请人支付{{请求金额}}。\n\n事实与理由：\n{{事实理由}}\n\n此致\n________人民法院\n\n申请人：________\n____年__月__日"),
                (12, "拍卖抵押物声请状", "诉讼类",
                 '[{"name":"声请人","hint":"例如：某银行","required":true},{"name":"相对人","hint":"例如：李某","required":true},{"name":"抵押物描述","hint":"例如：位于某某区XX路XX号的房产","required":true},{"name":"债权金额","hint":"例如：100万元","required":true}]',
                 "拍卖抵押物民事声请状\n\n声请人：{{声请人}}\n相对人：{{相对人}}\n\n为声请拍卖抵押物事：\n抵押物：{{抵押物描述}}\n债权金额：{{债权金额}}\n\n声请事项：请准予拍卖上开抵押物。\n\n事实及理由：相对人向声请人借款{{债权金额}}，并提供上开抵押物设定抵押权，因相对人逾期未还，声请人依法声请拍卖抵押物。\n\n此致\n________地方法院\n\n声请人：________\n____年__月__日"),
                (13, "抛弃继承声请状", "诉讼类",
                 '[{"name":"声请人","hint":"例如：张三","required":true},{"name":"被继承人","hint":"例如：张父","required":true},{"name":"关系","hint":"例如：长子","required":true}]',
                 "抛弃继承民事声请状\n\n声请人：{{声请人}}\n被继承人：{{被继承人}}\n\n声请人为被继承人之{{关系}}，因被继承人于____年__月__日死亡，声请人依法抛弃继承权。\n\n此致\n________地方法院\n\n声请人：________\n____年__月__日"),
                (14, "和解书", "其他文书",
                 '[{"name":"甲方","hint":"例如：张三","required":true},{"name":"乙方","hint":"例如：李四","required":true},{"name":"和解事由","hint":"例如：交通事故赔偿","required":true},{"name":"和解条件","hint":"例如：乙方一次性赔偿甲方医疗费、误工费等共计2万元，双方不再追究","required":true}]',
                 "和解书\n\n甲方：{{甲方}}\n乙方：{{乙方}}\n\n兹因{{和解事由}}，双方达成和解如下：\n一、{{和解条件}}\n二、甲方自愿放弃其他民事请求。\n三、双方就此事再无纠葛。\n\n本和解书一式两份，双方各执一份。\n\n甲方（签字）：________    乙方（签字）：________\n日期：____年__月__日"),
                (15, "消费争议申诉书", "其他文书",
                 '[{"name":"申诉人","hint":"例如：王小明","required":true},{"name":"被申诉企业","hint":"例如：某某电器商城","required":true},{"name":"争议事实","hint":"例如：购买的空调使用三天即损坏，商家拒绝退换","required":true},{"name":"请求事项","hint":"例如：要求退货并赔偿交通费100元","required":true}]',
                 "消费争议申诉书\n\n申诉人：{{申诉人}}\n被申诉企业：{{被申诉企业}}\n\n争议事实：\n{{争议事实}}\n\n请求事项：\n{{请求事项}}\n\n申诉人认为，根据《消费者权益保护法》，经营者应提供合格商品，对不合格商品应承担退换货责任。\n\n此致\n________消费者协会\n\n申诉人：________\n日期：____年__月__日"),
                (16, "存证信函", "其他文书",
                 '[{"name":"寄件人","hint":"例如：张三","required":true},{"name":"收件人","hint":"例如：李四","required":true},{"name":"主旨","hint":"例如：催告履行合同义务","required":true},{"name":"内容","hint":"例如：请贵方于收到本函后7日内交付货物，逾期将解除合同并追究违约责任","required":true}]',
                 "存证信函\n\n寄件人：{{寄件人}}\n收件人：{{收件人}}\n\n主旨：{{主旨}}\n\n说明：\n{{内容}}\n\n本函以存证信函方式寄送，作为法律证据。\n\n寄件人：________\n日期：____年__月__日"),
                (17, "自书遗嘱", "其他文书",
                 '[{"name":"立遗嘱人","hint":"例如：王德明","required":true},{"name":"财产分配","hint":"例如：名下房产由长子继承，存款由次女继承","required":true},{"name":"执行人","hint":"例如：长子王建国","required":true}]',
                 "自书遗嘱\n\n立遗嘱人：{{立遗嘱人}}，身份证号________。\n\n本人精神正常，自愿订立遗嘱如下：\n一、财产分配：{{财产分配}}\n二、指定{{执行人}}为遗嘱执行人。\n三、本遗嘱为本人最终意愿。\n\n立遗嘱人（签名）：________\n日期：____年__月__日"),
                (18, "国家损害赔偿请求书", "其他文书",
                 '[{"name":"请求人","hint":"例如：刘勇","required":true},{"name":"赔偿义务机关","hint":"例如：某某市公安局","required":true},{"name":"损害事实","hint":"例如：某年某月，被错误羁押30天","required":true},{"name":"请求金额","hint":"例如：10000元","required":true}]',
                 "国家损害赔偿请求书\n\n请求人：{{请求人}}\n赔偿义务机关：{{赔偿义务机关}}\n\n损害事实：\n{{损害事实}}\n\n请求赔偿金额：{{请求金额}}\n\n根据《国家赔偿法》，请求人依法提出赔偿申请。\n\n此致\n________赔偿义务机关\n\n请求人：________\n日期：____年__月__日"),
                (19, "交通违规异议书", "其他文书",
                 '[{"name":"异议人","hint":"例如：赵五","required":true},{"name":"违规单号","hint":"例如：京A123456","required":true},{"name":"异议理由","hint":"例如：当时信号灯为绿灯，不存在闯红灯行为","required":true}]',
                 "交通违规异议书\n\n异议人：{{异议人}}\n违规单号：{{违规单号}}\n\n异议理由：\n{{异议理由}}\n\n请贵单位查明事实，撤销该违规记录。\n\n此致\n________交通管理局\n\n异议人：________\n日期：____年__月__日"),
                (20, "学生申诉书", "其他文书",
                 '[{"name":"申诉人","hint":"例如：李小明","required":true},{"name":"被申诉单位","hint":"例如：某某大学教务处","required":true},{"name":"申诉事由","hint":"例如：因请假手续不全被记过处分，请求撤销","required":true}]',
                 "学生申诉书\n\n申诉人：{{申诉人}}\n被申诉单位：{{被申诉单位}}\n\n申诉事由：\n{{申诉事由}}\n\n申诉人认为，处分决定事实不清、程序不当，请求予以撤销。\n\n此致\n________大学学生申诉委员会\n\n申诉人：________\n日期：____年__月__日"),
                (21, "教师申诉书", "其他文书",
                 '[{"name":"申诉人","hint":"例如：张丽华","required":true},{"name":"被申诉单位","hint":"例如：某某市教育局","required":true},{"name":"申诉事由","hint":"例如：职称评审不公，请求复核","required":true}]',
                 "教师申诉书\n\n申诉人：{{申诉人}}\n被申诉单位：{{被申诉单位}}\n\n申诉事由：\n{{申诉事由}}\n\n申诉人请求依法复核职称评审结果，维护合法权益。\n\n此致\n________教育行政部门\n\n申诉人：________\n日期：____年__月__日"),
                (22, "请愿书", "其他文书",
                 '[{"name":"请愿人","hint":"例如：某某小区业主委员会","required":true},{"name":"请愿对象","hint":"例如：某某区政府","required":true},{"name":"请愿事项","hint":"例如：请求在小区门口增设公交站点","required":true}]',
                 "请愿书\n\n请愿人：{{请愿人}}\n请愿对象：{{请愿对象}}\n\n请愿事项：\n{{请愿事项}}\n\n请愿人代表广大居民，恳请贵单位予以解决。\n\n此致\n________人民政府\n\n请愿人代表：________\n日期：____年__月__日"),
                (23, "借据", "其他文书",
                 '[{"name":"贷与人","hint":"例如：陈先生","required":true},{"name":"借用人","hint":"例如：刘先生","required":true},{"name":"借款金额","hint":"例如：10万元","required":true},{"name":"还款期限","hint":"例如：2026年4月1日","required":true}]',
                 "借据\n\n贷与人：{{贷与人}}\n借用人：{{借用人}}\n\n兹借到人民币{{借款金额}}，并约定于{{还款期限}}前归还。逾期未还，每日按借款金额的0.05%支付违约金。\n\n恐口无凭，特立此据。\n\n立据人：________\n日期：____年__月__日"),
            ]
            c.executemany("INSERT INTO templates VALUES (?,?,?,?,?)", templates)
        conn.commit()
        conn.close()
    except Exception as e:
        print("数据库初始化失败:", e)

init_db()

# ---------- 配置文件（记住密码）----------
def save_login_credentials(username, password, remember):
    data = {"username": username, "password": password, "remember": remember}
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f)

def load_login_credentials():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if data.get("remember", False):
                return data.get("username", ""), data.get("password", ""), True
    return "", "", False

# ---------- 数据库操作函数 ----------
def register_user(phone, username, password):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO users (phone, username, password) VALUES (?,?,?)", (phone, username, password))
        conn.commit()
        return True
    except:
        return False
    finally:
        conn.close()

def login_user(username, password):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, is_purchased FROM users WHERE username=? AND password=?", (username, password))
    user = c.fetchone()
    conn.close()
    return user

def update_purchase(user_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE users SET is_purchased=1 WHERE id=?", (user_id,))
    conn.commit()
    conn.close()

def get_templates_by_category(category):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, name, fields, content_template FROM templates WHERE category=? ORDER BY name", (category,))
    rows = c.fetchall()
    templates = []
    for row in rows:
        try:
            fields = json.loads(row[2])
        except:
            fields = []
        templates.append({
            "id": row[0],
            "name": row[1],
            "fields": fields,
            "content_template": row[3]
        })
    conn.close()
    return templates

def save_document(user_id, template_id, filled_data, generated_content):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("INSERT INTO documents (user_id, template_id, filled_data, generated_content, create_time) VALUES (?,?,?,?,?)",
              (user_id, template_id, json.dumps(filled_data), generated_content, datetime.now().isoformat()))
    conn.commit()
    conn.close()

def get_user_documents(user_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, template_id, generated_content, create_time FROM documents WHERE user_id=? ORDER BY create_time DESC", (user_id,))
    rows = c.fetchall()
    docs = []
    for row in rows:
        c.execute("SELECT name FROM templates WHERE id=?", (row[1],))
        tname = c.fetchone()
        docs.append({
            "doc_id": row[0],
            "template_name": tname[0] if tname else "未知模板",
            "content": row[2],
            "create_time": row[3]
        })
    conn.close()
    return docs

def reset_password_by_phone(phone, new_password):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE users SET password=? WHERE phone=?", (new_password, phone))
    if c.rowcount > 0:
        conn.commit()
        conn.close()
        return True
    conn.close()
    return False

# ---------- 试用期管理 ----------
def init_trial():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT install_date, last_run_date FROM trial LIMIT 1")
    row = c.fetchone()
    now = datetime.now()
    now_str = now.strftime("%Y-%m-%d %H:%M:%S")

    if row is None:
        c.execute("INSERT INTO trial (install_date, last_run_date) VALUES (?,?)", (now_str, now_str))
        conn.commit()
        conn.close()
        return True, 2 * 24 * 3600, False
    else:
        install_date_str, last_run_date_str = row
        install_date = datetime.strptime(install_date_str, "%Y-%m-%d %H:%M:%S")
        last_run_date = datetime.strptime(last_run_date_str, "%Y-%m-%d %H:%M:%S")
        if now < last_run_date:  # 时间回拨
            conn.close()
            return False, 0, True
        c.execute("UPDATE trial SET last_run_date = ?", (now_str,))
        conn.commit()
        conn.close()
        elapsed = (now - install_date).total_seconds()
        trial_secs = 2 * 24 * 3600
        if elapsed >= trial_secs:
            return False, 0, False
        else:
            return True, trial_secs - elapsed, False

# ---------- 激活码 ----------
UNLOCK_CODE = "LEGAL2025"

def show_purchase_dialog(page, on_success):
    def close_dlg(e):
        dlg.open = False
        page.update()

    def verify(e):
        if code_input.value == UNLOCK_CODE:
            dlg.open = False
            page.update()
            on_success()
        else:
            error_label.value = "激活码错误"
            page.update()

    code_input = TextField(label="激活码", hint_text="请输入激活码")
    error_label = Text("", color=ft.Colors.RED)
    dlg = AlertDialog(
        title=Text("解锁全部模板"),
        content=Column([
            Text("请使用微信/支付宝扫码支付 ¥19.90"),
            Text("付款后联系客服获取激活码"),
            Text("客服微信：SummerDeun999(备注：法律文书)"),
            code_input,
            error_label
        ], tight=True),
        actions=[
            TextButton("取消", on_click=close_dlg),
            ElevatedButton("验证并解锁", on_click=verify)
        ]
    )
    page.dialog = dlg
    dlg.open = True
    page.update()

# ---------- Flet 主函数 ----------
def main(page: ft.Page):
    page.title = "法律文书助手"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    page.vertical_alignment = ft.MainAxisAlignment.START
    page.window_width = 400
    page.window_height = 700

    current_user_id = None
    current_user_purchased = False

    countdown_text = ft.Text("", size=16, color=ft.Colors.ORANGE)

    async def update_countdown():
        while True:
            valid, remaining, tampered = init_trial()
            if tampered:
                countdown_text.value = "系统时间异常，请联系客服"
                countdown_text.color = ft.Colors.RED
            elif not valid:
                countdown_text.value = "试用期已结束，请付费解锁"
                countdown_text.color = ft.Colors.RED
            else:
                days = int(remaining // 86400)
                hours = int((remaining % 86400) // 3600)
                mins = int((remaining % 3600) // 60)
                secs = int(remaining % 60)
                countdown_text.value = f"免费试用剩余：{days}天 {hours:02d}时 {mins:02d}分 {secs:02d}秒"
            page.update()
            await asyncio.sleep(1)

    page.run_task(update_countdown)

    def route_change(e):
        page.views.clear()
        if page.route == "/":
            page.views.append(login_view())
        elif page.route == "/register":
            page.views.append(register_view())
        elif page.route == "/home":
            page.views.append(home_view())
        elif page.route == "/template":
            template_id = int(page.route.split("?id=")[1])
            page.views.append(template_form_view(template_id))
        elif page.route == "/preview":
            page.views.append(preview_view())
        elif page.route == "/profile":
            page.views.append(profile_view())
        elif page.route == "/history":
            page.views.append(history_view())
        elif page.route == "/view_doc":
            doc_id = int(page.route.split("?id=")[1])
            page.views.append(view_doc_view(doc_id))
        page.update()

    page.on_route_change = route_change
    page.go("/")

    def login_view():
        username_input = TextField(label="用户名", autofocus=True)
        pwd_input = TextField(label="密码", password=True, can_reveal_password=True)
        remember_check = ft.Checkbox(label="记住密码", value=False)
        error_text = Text("", color=ft.Colors.RED)

        saved_user, saved_pwd, remember = load_login_credentials()
        if saved_user:
            username_input.value = saved_user
            pwd_input.value = saved_pwd
            remember_check.value = remember

        def do_login(e):
            nonlocal current_user_id, current_user_purchased
            user = login_user(username_input.value, pwd_input.value)
            if user:
                current_user_id = user[0]
                current_user_purchased = bool(user[1])
                save_login_credentials(username_input.value, pwd_input.value, remember_check.value)
                page.go("/home")
            else:
                error_text.value = "用户名或密码错误"
                page.update()

        def forgot_password(e):
            def reset(e):
                phone = phone_input.value
                new_pwd = new_pwd_input.value
                confirm = confirm_input.value
                if not phone or not new_pwd:
                    error_msg.value = "手机号和新密码不能为空"
                    dlg.update()
                    return
                if new_pwd != confirm:
                    error_msg.value = "两次输入不一致"
                    dlg.update()
                    return
                if reset_password_by_phone(phone, new_pwd):
                    dlg.open = False
                    page.update()
                    page.snack_bar = ft.SnackBar(content=Text("密码重置成功，请重新登录"))
                    page.snack_bar.open = True
                    page.update()
                else:
                    error_msg.value = "未找到该手机号"
                    dlg.update()

            phone_input = TextField(label="注册手机号")
            new_pwd_input = TextField(label="新密码", password=True)
            confirm_input = TextField(label="确认新密码", password=True)
            error_msg = Text("", color=ft.Colors.RED)
            dlg = AlertDialog(
                title=Text("找回密码"),
                content=Column([phone_input, new_pwd_input, confirm_input, error_msg], tight=True),
                actions=[TextButton("取消", on_click=lambda e: setattr(dlg, 'open', False)), ElevatedButton("重置", on_click=reset)]
            )
            page.dialog = dlg
            dlg.open = True
            page.update()

        return ft.View("/", [
            Container(
                width=300,
                content=Column(
                    spacing=20,
                    controls=[
                        Text("法律文书助手", size=30, weight=ft.FontWeight.BOLD),
                        username_input,
                        pwd_input,
                        remember_check,
                        error_text,
                        ElevatedButton("登录", on_click=do_login),
                        Row([
                            TextButton("注册新账号", on_click=lambda e: page.go("/register")),
                            TextButton("忘记密码？", on_click=forgot_password),
                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                    ]
                )
            )
        ])

    def register_view():
        phone_input = TextField(label="手机号", keyboard_type=ft.KeyboardType.PHONE)
        username_input = TextField(label="用户名")
        pwd_input = TextField(label="密码", password=True)
        confirm_input = TextField(label="确认密码", password=True)
        error_text = Text("", color=ft.Colors.RED)

        def do_register(e):
            if pwd_input.value != confirm_input.value:
                error_text.value = "两次密码不一致"
                page.update()
                return
            if register_user(phone_input.value, username_input.value, pwd_input.value):
                page.go("/")
            else:
                error_text.value = "用户名或手机号已存在"
                page.update()

        return ft.View("/register", [
            Container(
                width=300,
                content=Column(
                    spacing=20,
                    controls=[
                        Text("注册账号", size=30),
                        phone_input,
                        username_input,
                        pwd_input,
                        confirm_input,
                        error_text,
                        ElevatedButton("注册", on_click=do_register),
                        TextButton("返回登录", on_click=lambda e: page.go("/")),
                    ]
                )
            )
        ])

    def home_view():
        def show_category(category):
            nonlocal current_user_purchased
            if not current_user_purchased:
                valid, _, tampered = init_trial()
                if tampered:
                    page.snack_bar = ft.SnackBar(content=Text("系统时间异常，请联系客服"))
                    page.snack_bar.open = True
                    page.update()
                    return
                if not valid:
                    def after_unlock():
                        nonlocal current_user_purchased
                        update_purchase(current_user_id)
                        current_user_purchased = True
                        show_category(category)
                    reply = ft.AlertDialog(
                        title=Text("试用期已过"),
                        content=Text(f"“{category}”中的模板需要付费后才能使用。\n是否立即解锁全部模板？"),
                        actions=[
                            TextButton("取消", on_click=lambda e: setattr(reply, 'open', False)),
                            ElevatedButton("解锁", on_click=lambda e: show_purchase_dialog(page, after_unlock))
                        ]
                    )
                    page.dialog = reply
                    reply.open = True
                    page.update()
                    return
            templates = get_templates_by_category(category)
            if not templates:
                page.snack_bar = ft.SnackBar(content=Text(f"暂无{category}模板"))
                page.snack_bar.open = True
                page.update()
                return
            def open_template(t):
                page.go(f"/template?id={t['id']}")
            items = [ElevatedButton(t["name"], on_click=lambda e, t=t: open_template(t)) for t in templates]
            dlg = AlertDialog(
                title=Text(f"{category} - 选择模板"),
                content=Column(items, scroll=ft.ScrollMode.AUTO),
                actions=[TextButton("关闭", on_click=lambda e: setattr(dlg, 'open', False))]
            )
            page.dialog = dlg
            dlg.open = True
            page.update()

        return ft.View("/home", [
            AppBar(title=Text("法律文书模板"), center_title=True, actions=[
                IconButton(Icons.PERSON, on_click=lambda e: page.go("/profile"))
            ]),
            Container(
                content=Column([
                    countdown_text,
                    ElevatedButton("合同类", on_click=lambda e: show_category("合同类"), width=200, height=60),
                    ElevatedButton("诉讼类", on_click=lambda e: show_category("诉讼类"), width=200, height=60),
                    ElevatedButton("其他文书", on_click=lambda e: show_category("其他文书"), width=200, height=60),
                ], spacing=30, horizontal_alignment=ft.CrossAxisAlignment.CENTER),
                expand=True,
                alignment=ft.alignment.center
            )
        ])

    def template_form_view(template_id):
        # 获取所有模板
        all_templates = []
        for cat in ["合同类", "诉讼类", "其他文书"]:
            all_templates.extend(get_templates_by_category(cat))
        template = next((t for t in all_templates if t["id"] == template_id), None)
        if not template:
            return ft.View("/template", [Text("模板不存在")])

        field_inputs = {}
        for field in template["fields"]:
            field_inputs[field["name"]] = TextField(
                label=field["name"],
                hint_text=field["hint"],
                multiline=field["name"] in ["诉讼请求", "事实与理由", "子女抚养安排", "财产分割请求", "上诉理由", "事实与证据", "答辩事由", "事实理由", "争议事实", "内容", "财产分配", "损害事实", "异议理由", "申诉事由", "请愿事项", "和解条件"]
            )

        def generate_doc(e):
            filled = {name: w.value or "" for name, w in field_inputs.items()}
            content = template["content_template"]
            for key, val in filled.items():
                content = content.replace("{{"+key+"}}", val)
            save_document(current_user_id, template_id, filled, content)
            page.client_storage.set("preview_content", content)
            page.client_storage.set("preview_title", template["name"])
            page.go("/preview")

        return ft.View(f"/template?id={template_id}", [
            AppBar(title=Text(template["name"]), leading=IconButton(Icons.ARROW_BACK, on_click=lambda e: page.go("/home"))),
            ListView(
                controls=[
                    Column(
                        spacing=20,
                        controls=[
                            Text("请填写以下信息：", size=18, weight=ft.FontWeight.BOLD),
                            *field_inputs.values(),
                            ElevatedButton("生成文书", on_click=generate_doc)
                        ]
                    )
                ],
                expand=True,
                padding=20
            )
        ])

    def preview_view():
        content = page.client_storage.get("preview_content") or ""
        title = page.client_storage.get("preview_title") or "法律文书"

        def export_word(e):
            doc = Document()
            section = doc.sections[0]
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            for line in content.split("\n"):
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Cm(0.75)
                p.paragraph_format.line_spacing = 1.5
                p.style.font.size = Pt(12)
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{title}_{current_user_id}.docx")
            doc.save(file_path)
            webbrowser.open(file_path)
            page.snack_bar = ft.SnackBar(content=Text("已生成Word文档，保存在临时文件夹"), duration=2000)
            page.snack_bar.open = True
            page.update()

        def print_doc(e):
            html = f"""
            <!DOCTYPE html>
            <html>
            <head><meta charset="UTF-8"><title>{title}</title>
            <style>
                @page {{ size: A4; margin: 2.54cm; }}
                body {{ font-family: '宋体', SimSun; font-size: 12pt; line-height: 1.5; }}
                h1 {{ text-align: center; font-size: 18pt; }}
                .content {{ white-space: pre-wrap; text-indent: 2em; }}
            </style>
            </head>
            <body>
                <h1>{title}</h1>
                <div class="content">{content.replace(chr(10), '<br>')}</div>
                <script>window.print();</script>
            </body>
            </html>
            """
            temp_html = os.path.join(tempfile.gettempdir(), "print_preview.html")
            with open(temp_html, "w", encoding="utf-8") as f:
                f.write(html)
            webbrowser.open(temp_html)

        return ft.View("/preview", [
            AppBar(title=Text("文书预览"), leading=IconButton(Icons.ARROW_BACK, on_click=lambda e: page.go("/home"))),
            Container(
                content=Column([
                    Text(content, selectable=True, size=16),
                    Row([
                        ElevatedButton("导出Word", icon=Icons.SAVE, on_click=export_word),
                        ElevatedButton("打印", icon=Icons.PRINT, on_click=print_doc),
                    ], alignment=ft.MainAxisAlignment.CENTER)
                ], scroll=ft.ScrollMode.AUTO),
                padding=20,
                expand=True
            )
        ])

    def profile_view():
        docs = get_user_documents(current_user_id)
        doc_items = []
        for doc in docs:
            time_str = doc["create_time"][:19].replace("T", " ")
            item = Container(
                margin=5,
                padding=10,
                border=ft.border.all(color=ft.Colors.GREY_300),
                border_radius=10,
                content=Column([
                    Text(doc["template_name"], size=16, weight=ft.FontWeight.BOLD),
                    Text(f"生成时间：{time_str}", size=12, color=ft.Colors.GREY),
                    ElevatedButton("查看详情", on_click=lambda e, did=doc["doc_id"]: page.go(f"/view_doc?id={did}"))
                ])
            )
            doc_items.append(item)
        if not doc_items:
            doc_items.append(Text("暂无历史记录", size=16, color=ft.Colors.GREY))

        col_controls = [
            Text(f"已登录：用户ID {current_user_id}", size=16),
            Text(f"付费状态：{'已解锁' if current_user_purchased else '未解锁'}"),
            Divider(),
            Text("我的文书", size=18, weight=ft.FontWeight.BOLD),
            Column(doc_items, scroll=ft.ScrollMode.AUTO, height=300),
            ElevatedButton("退出登录", on_click=lambda e: page.go("/"))
        ]
        if not current_user_purchased:
            col_controls.insert(2, ElevatedButton("解锁全部模板", on_click=lambda e: show_purchase_dialog(page, lambda: after_unlock())))

        def after_unlock():
            nonlocal current_user_purchased
            update_purchase(current_user_id)
            current_user_purchased = True
            page.go("/profile")

        return ft.View("/profile", [
            AppBar(title=Text("个人中心"), leading=IconButton(Icons.ARROW_BACK, on_click=lambda e: page.go("/home"))),
            Column(col_controls, spacing=20, horizontal_alignment=ft.CrossAxisAlignment.CENTER)
        ])

    def history_view():
        return ft.View("/history", [Text("历史记录已整合到个人中心", size=20)])

    def view_doc_view(doc_id):
        docs = get_user_documents(current_user_id)
        doc = next((d for d in docs if d["doc_id"] == doc_id), None)
        if not doc:
            return ft.View("/view_doc", [Text("文书不存在")])
        content = doc["content"]
        title = doc["template_name"]

        def export_word(e):
            doc_obj = Document()
            section = doc_obj.sections[0]
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            title_para = doc_obj.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style = doc_obj.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            for line in content.split("\n"):
                p = doc_obj.add_paragraph(line)
                p.paragraph_format.first_line_indent = Cm(0.75)
                p.paragraph_format.line_spacing = 1.5
                p.style.font.size = Pt(12)
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"{title}_{current_user_id}_{doc_id}.docx")
            doc_obj.save(file_path)
            webbrowser.open(file_path)
            page.snack_bar = ft.SnackBar(content=Text("已生成Word文档"), duration=2000)
            page.snack_bar.open = True
            page.update()

        def print_doc(e):
            html = f"""
            <!DOCTYPE html>
            <html>
            <head><meta charset="UTF-8"><title>{title}</title>
            <style>
                @page {{ size: A4; margin: 2.54cm; }}
                body {{ font-family: '宋体', SimSun; font-size: 12pt; line-height: 1.5; }}
                h1 {{ text-align: center; font-size: 18pt; }}
                .content {{ white-space: pre-wrap; text-indent: 2em; }}
            </style>
            </head>
            <body>
                <h1>{title}</h1>
                <div class="content">{content.replace(chr(10), '<br>')}</div>
                <script>window.print();</script>
            </body>
            </html>
            """
            temp_html = os.path.join(tempfile.gettempdir(), "print_preview.html")
            with open(temp_html, "w", encoding="utf-8") as f:
                f.write(html)
            webbrowser.open(temp_html)

        return ft.View(f"/view_doc?id={doc_id}", [
            AppBar(title=Text("文书详情"), leading=IconButton(Icons.ARROW_BACK, on_click=lambda e: page.go("/profile"))),
            Container(
                content=Column([
                    Text(content, selectable=True, size=16),
                    Row([
                        ElevatedButton("导出Word", icon=Icons.SAVE, on_click=export_word),
                        ElevatedButton("打印", icon=Icons.PRINT, on_click=print_doc),
                    ], alignment=ft.MainAxisAlignment.CENTER)
                ], scroll=ft.ScrollMode.AUTO),
                padding=20,
                expand=True
            )
        ])

    page.update()

if __name__ == "__main__":
    ft.app(target=main)